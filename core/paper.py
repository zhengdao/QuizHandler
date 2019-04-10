#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
core.paper
~~~~~~~~~~~~~~~~~~

This module define the questions paper.

"""

from core import nls
from core.quiz import Config
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


class QuizPaper:
    """
    Define the QuizPaper.
    """

    def __init__(self, title='', questions=None):
        self.title = title

        if questions is None:
            self.__questions = dict()
        else:
            self.__questions = questions

    def settitle(self, title):
        self.title = title

    def setquestions(self, questions):
        if isinstance(questions, dict):
            self.__questions = questions

    def addquestion(self, category, question):
        qset = self.get_qset_by_type(category)
        qset.append(question)

    def get_qset_by_type(self, category):
        qset = self.__questions.get(category)
        if qset is None:
            self.__questions[category] = []

        return self.__questions[category]

    @staticmethod
    def preparestyles(doc):
        ustyles = doc.styles

        # Customize the default font
        style = ustyles['Normal']
        style.font.size = Pt(12)
        style.font.name = u'宋体'
        style.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

        pformat = style.paragraph_format
        pformat.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        pformat.space_before = Pt(0)
        pformat.space_after = Pt(0)
        pformat.line_spacing_rule = WD_LINE_SPACING.SINGLE

        # Define paper title styles
        style = ustyles.add_style('title', WD_STYLE_TYPE.PARAGRAPH)
        pfont = style.font
        pfont.size = Pt(16)
        pfont.bold = True
        pfont.name = u'宋体'
        style.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

        pformat = style.paragraph_format
        pformat.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        pformat.space_before = Pt(0)
        pformat.space_after = Pt(16)

        # Define paragraph styles
        style = ustyles.add_style('pstyle', WD_STYLE_TYPE.PARAGRAPH)
        pfont = style.font
        pfont.bold = None
        pfont.name = u'宋体'
        style.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

        style = ustyles.add_style('ptitle', WD_STYLE_TYPE.CHARACTER)
        pfont = style.font
        pfont.bold = True
        pfont.name = u'宋体'
        style.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

        style = ustyles.add_style('pcontent', WD_STYLE_TYPE.CHARACTER)
        pfont = style.font
        pfont.bold = None
        pfont.name = u'宋体'
        style.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

        # Customize "Header" style
        hstyle = ustyles['Heading 1']
        style = ustyles.add_style('h1', WD_STYLE_TYPE.PARAGRAPH)
        pfont = style.font
        pfont.bold = True
        pfont.size = Pt(13)  # hstyle.font.size
        pfont.name = u'宋体'
        style.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        pfont.color.rgb = hstyle.font.color.rgb

        hstyle = ustyles['Heading 2']
        style = ustyles.add_style('h2', WD_STYLE_TYPE.PARAGRAPH)
        pfont = style.font
        pfont.bold = True
        pfont.size = Pt(13)  # hstyle.font.size
        style.font.name = u'宋体'
        style.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        pfont.color.rgb = RGBColor(0, 0x20, 0x60)  # hstyle.font.color.rgb

        pformat = style.paragraph_format
        pformat.space_before = Pt(12)
        pformat.space_after = Pt(6)

    def toword(self, file, config=Config()):
        doc = Document()

        # prepare styles
        QuizPaper.preparestyles(doc)

        # Generate top title
        doc.add_paragraph(self.title, 'title')

        # Generate questions
        i = 0
        qdict = self.__questions
        for (category, qset) in qdict.items():
            # Generate question category
            txt = self.getbulletlead(i) \
                  + nls.getNLSText('N-' + category)
            # doc.add_heading(txt, 1)
            doc.add_paragraph(txt, 'h1')

            # Generate quests
            j = 0
            for j in range(len(qset)):
                question = qset[j]
                question.toword(doc, self.getstemlead(j), 2, config)

            i = i + 1

        doc.save(file)

    linums = [
        u'一', u'二', u'三', u'四', u'五', u'六',
        u'七', u'八', u'九', u'十'
    ]

    def getbulletlead(self, i=0, category=None):
        return self.linums[i] + u'、' + ''

    def getstemlead(self, i=0):
        return str(i + 1) + u'.' + ''

    def __str__(self):
        tmp = [self.title, '\n\n']

        cnt = 0
        i = 0
        qdict = self.__questions
        for (category, qset) in qdict.items():
            tmp.extend(
                [self.getbulletlead(i), nls.getNLSText('N-' + category),
                 '\n'])

            for j in range(len(qset)):
                cnt = cnt + 1
                question = qset[j]
                tmp.extend([self.getstemlead(j), str(question), '\n'])
            tmp.append('\n')

            i = i + 1

        tmp.extend([nls.getNLSText('msgQuizNum') + str(cnt)])

        return ''.join(tmp)

    __repr__ = __str__
