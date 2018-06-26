#!/usr/bin/env python
# -*- coding: utf-8 -*-

from dep import NLSFactory
from quiz.Config import Config
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING


class QuizPaper:

    def __init__(self, title='', questions=dict()):
        self.title = title
        self.questions = questions

    def setTitle(self, title):
        self.title = title

    def setQuestions(self, questions):
        if isinstance(questions, dict):
            self.questions = questions

    def addQuestion(self, category, question):
        qset = self.getQuestionSetByType(category)
        qset.append(question)

    def getQuestionSetByType(self, category):
        qset = self.questions.get(category)
        if qset is None:
            self.questions[category] = []

        return self.questions[category]

    def prepareStyles(self, doc):
        ustyles = doc.styles

        # Customize the default font
        style = ustyles['Normal']
        style.font.size = Pt(12)
        style.font.name = u'宋体'
        style.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

        pformat = style.paragraph_format
        pformat.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        pformat.space_before = Pt(0)
        pformat.space_after = Pt(0)
        pformat.line_spacing_rule = WD_LINE_SPACING.SINGLE

        # Define paper title styles
        style = ustyles.add_style('title', WD_STYLE_TYPE.PARAGRAPH)
        pfont = style.font
        pfont.size = Pt(18)
        pfont.bold = True
        pfont.name = u'宋体'
        style.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

        pformat = style.paragraph_format
        pformat.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        pformat.space_before = Pt(4)
        pformat.space_after = Pt(12)

        # Define paragraph styles
        style = ustyles.add_style('pstyle', WD_STYLE_TYPE.PARAGRAPH)
        pfont = style.font
        pfont.bold = None
        pfont.name = u'宋体'
        style.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

        style = ustyles.add_style('ptitle', WD_STYLE_TYPE.CHARACTER)
        pfont = style.font
        pfont.bold = True
        pfont.name = u'宋体'
        style.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

        style = ustyles.add_style('pcontent', WD_STYLE_TYPE.CHARACTER)
        pfont = style.font
        pfont.bold = None
        pfont.name = u'宋体'
        style.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

        # Customize "Header" style
        hstyle = ustyles['Heading 1']
        style = ustyles.add_style('h1', WD_STYLE_TYPE.PARAGRAPH)
        pfont = style.font
        pfont.bold = True
        pfont.size = Pt(13)  # hstyle.font.size
        pfont.name = u'宋体'
        style.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        pfont.color.rgb = hstyle.font.color.rgb

        hstyle = ustyles['Heading 2']
        style = ustyles.add_style('h2', WD_STYLE_TYPE.PARAGRAPH)
        pfont = style.font
        pfont.bold = True
        pfont.size = Pt(13)  # hstyle.font.size
        style.font.name = u'宋体'
        style.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        pfont.color.rgb = RGBColor(0, 0x20, 0x60)  # hstyle.font.color.rgb

        pformat = style.paragraph_format
        pformat.space_before = Pt(12)
        pformat.space_after = Pt(6)

    def toWord(self, file, config=Config()):
        doc = Document()

        # prepare styles
        self.prepareStyles(doc)

        # Generate top title
        doc.add_paragraph(self.title, 'title')

        # Generate questions
        i = 0
        qdict = self.questions
        for (category, qset) in qdict.items():
            # Generate question category
            txt = self.getBulletLead(i) \
                  + NLSFactory.getNLSText('N-' + category)
            # doc.add_heading(txt, 1)
            doc.add_paragraph(txt, 'h1')

            # Generate quests
            j = 0
            for j in range(len(qset)):
                question = qset[j]
                question.toWord(doc, self.getStemLead(j), 2, config)

            i = i + 1

        doc.save(file)

    linums = [
        u'一', u'二', u'三', u'四', u'五', u'六',
        u'七', u'八', u'九', u'十'
    ]

    def getBulletLead(self, i=0, category=None):
        return self.linums[i] + u'、' + ''

    def getStemLead(self, i=0):
        return str(i + 1) + u'.' + ''

    def __str__(self):
        tmp = [self.title, '\n']

        cnt = 0
        i = 0
        qdict = self.questions
        for (category, qset) in qdict.items():
            tmp.extend(
                [self.getBulletLead(i), NLSFactory.getNLSText('N-' + category),
                 '\n'])

            j = 0
            for j in range(len(qset)):
                cnt = cnt + 1
                question = qset[j]
                tmp.extend([self.getStemLead(j), str(question), '\n'])
            tmp.append('\n')

            i = i + 1

        tmp.extend([NLSFactory.getNLSText('msgQuizNum') + str(cnt)])

        return ''.join(tmp)

    __repr__ = __str__
